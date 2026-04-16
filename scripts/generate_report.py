"""
generate_report.py — Genera el informe técnico del Hito 3 en formato Word (.docx).

Uso:
    uv run python scripts/generate_report.py
    # o
    python scripts/generate_report.py

Produce: Hito3_Informe_Tecnico.docx en la raíz del proyecto.
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# ──────────────────────────────────────────────────────────────
# Constantes
# ──────────────────────────────────────────────────────────────

OUTPUT_PATH = Path(__file__).resolve().parents[1] / "Hito3_Informe_Tecnico.docx"

TEAM = [
    "Juan Jacobo Delgado",
    "Gabriel Martinez",
    "Nicolas Cuaran",
    "Juan Jose Orozco",
    "Sebastian Belalcazar",
]

COLOR_TITULO = RGBColor(0x1F, 0x4E, 0x79)   # azul oscuro
COLOR_SECCION = RGBColor(0x2E, 0x75, 0xB6)  # azul medio
COLOR_ACCENT = RGBColor(0xC0, 0x00, 0x00)   # rojo


# ──────────────────────────────────────────────────────────────
# Utilidades de formato
# ──────────────────────────────────────────────────────────────

def set_heading_color(paragraph, color: RGBColor) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = color


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table_header(table, headers: list[str], color: RGBColor = COLOR_SECCION) -> None:
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Color de fondo de la celda
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        tcPr.append(shd)


def fill_table_row(row, values: list[str]) -> None:
    for i, val in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = val


# ──────────────────────────────────────────────────────────────
# Secciones del documento
# ──────────────────────────────────────────────────────────────

def build_portada(doc: Document) -> None:
    # Espacio inicial
    for _ in range(4):
        doc.add_paragraph()

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SEMINARIO INGENIERÍA DE DATOS E IA")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITULO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Universidad Autónoma de Occidente — UAO")
    run2.font.size = Pt(14)
    run2.font.color.rgb = COLOR_SECCION

    for _ in range(2):
        doc.add_paragraph()

    # Título del hito
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        "HITO 3: Metodología de Evaluación Causal\n"
        "y Primeros Resultados con Incertidumbre"
    )
    run3.font.size = Pt(18)
    run3.font.bold = True
    run3.font.color.rgb = COLOR_TITULO

    for _ in range(2):
        doc.add_paragraph()

    # Subtítulo del proyecto
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(
        "Evaluación de Políticas de Acceso a Educación Superior\n"
        "Gobierno Nacional Colombia 2022–2026 (Gobierno Petro)"
    )
    run4.font.size = Pt(13)
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    for _ in range(3):
        doc.add_paragraph()

    # Equipo
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Equipo de trabajo:")
    run5.font.bold = True
    run5.font.size = Pt(12)

    for member in TEAM:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.add_run(member).font.size = Pt(11)

    for _ in range(2):
        doc.add_paragraph()

    # Fecha
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run(f"Bogotá D.C., {datetime.now().strftime('%B %Y').title()}")
    run6.font.size = Pt(11)
    run6.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()


def build_resumen_ejecutivo(doc: Document) -> None:
    h = doc.add_heading("Resumen Ejecutivo", level=1)
    set_heading_color(h, COLOR_TITULO)

    doc.add_paragraph(
        "El presente informe documenta el desarrollo del Hito 3 del Seminario de "
        "Ingeniería de Datos e IA de la Universidad Autónoma de Occidente. Este hito "
        "corresponde a la fase de metodología de evaluación causal y primeros resultados "
        "con análisis de incertidumbre, en el marco del proyecto de evaluación de las "
        "políticas educativas del Gobierno Petro (2022–2026)."
    )

    doc.add_paragraph(
        "El trabajo consolidado en este hito abarca:"
    )

    items = [
        "Pipeline ETL completo (Hito 2): descarga de datos SNIES, ICFES y PND desde "
        "Google Drive, carga en PostgreSQL con esquema estrella (raw → unified → facts), "
        "con 11 pasos secuenciales, calidad de datos y exportación automática.",
        "Metodología de evaluación causal (Criterio D): implementación de Series de "
        "Tiempo Interrumpidas (ITS) y Diferencias en Diferencias (DiD) para estimar el "
        "efecto diferencial de la política en IES Oficiales vs. Privadas.",
        "Validación e incertidumbre (Criterio F): bootstrap por bloques (N=1,000 "
        "remuestras), pruebas placebo temporales y sectoriales, análisis de sensibilidad "
        "al punto de quiebre y a la composición de la muestra.",
        "Dashboard Streamlit interactivo con 5 pestañas: Tendencias, ITS, DiD, Bootstrap "
        "y Resumen ejecutivo.",
        "Notebook Jupyter documentado con ejecución paso a paso del análisis.",
        "Este informe técnico como entregable documentado del hito.",
    ]

    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph(
        "Los resultados se presentan bajo principios de neutralidad técnica: los "
        "estimadores son medidas de asociación (contribución), no de causalidad pura. "
        "Ningún hallazgo se interpreta como validación ni invalidación de una posición política."
    )
    doc.add_page_break()


def build_introduccion(doc: Document) -> None:
    h = doc.add_heading("1. Introducción y Contexto", level=1)
    set_heading_color(h, COLOR_TITULO)

    h2 = doc.add_heading("1.1 Pregunta de evaluación central", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "El proyecto busca responder la siguiente pregunta de política pública:"
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "¿En qué medida las políticas de acceso a educación superior del Gobierno Petro "
        "(2022–2026) —en particular la gratuidad y los programas de matrícula— modificaron "
        "la matrícula y la primera matrícula en Instituciones de Educación Superior (IES) "
        "públicas, respecto a la tendencia previa y al comportamiento de las IES privadas "
        "como grupo de control?"
    )
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)

    doc.add_paragraph(
        "Esta pregunta operacionaliza las preguntas P1 y P6 del Hito 1 sobre acceso a "
        "educación superior, y se alinea con el indicador SINERGIA ID 91 ('Estudiantes "
        "nuevos en Educación Superior Pública') del Plan Nacional de Desarrollo 2022–2026."
    )

    h2 = doc.add_heading("1.2 Marco de identificación causal", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "La política educativa del Gobierno Petro opera a nivel nacional sin variación "
        "aleatoria en la asignación de tratamiento. Esto implica que:"
    )

    for txt in [
        "No existe un grupo de control 'puro' en educación superior pública (toda IES "
        "oficial se ve afectada).",
        "La atribución completa es inviable sin un experimento natural o discontinuidad "
        "geográfica clara.",
        "Se aplica el criterio de contribución (no atribución): estimar en qué dirección "
        "y magnitud la política está correlacionada con cambios observados, controlando "
        "tendencias preexistentes y covariables de contexto.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    h2 = doc.add_heading("1.3 Estrategia de identificación adoptada", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "Se combinan dos enfoques complementarios para triangulación:"
    )

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    add_table_header(table, ["Método", "Pregunta que responde", "Supuesto clave"])
    fill_table_row(
        table.rows[1],
        [
            "ITS (Series de Tiempo Interrumpidas)",
            "¿Cambió el nivel o la tendencia de matrícula en IES públicas al inicio del gobierno?",
            "La tendencia pre-2022 habría continuado sin la política",
        ],
    )
    fill_table_row(
        table.rows[2],
        [
            "DiD (Diferencias en Diferencias)",
            "¿Cambió la matrícula en IES públicas más que en privadas tras 2022?",
            "Tendencias paralelas pre-tratamiento entre sectores",
        ],
    )

    doc.add_paragraph()
    doc.add_page_break()


def build_fuentes_datos(doc: Document) -> None:
    h = doc.add_heading("2. Fuentes de Datos", level=1)
    set_heading_color(h, COLOR_TITULO)

    doc.add_paragraph(
        "El proyecto utiliza tres fuentes de datos oficiales descargadas desde una "
        "carpeta compartida de Google Drive:"
    )

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    add_table_header(table, ["Fuente", "Categorías", "Periodo", "Formato"])
    fill_table_row(
        table.rows[1],
        [
            "SNIES (Sistema Nacional de Información de Educación Superior)",
            "Administrativos, admitidos, docentes, graduados, inscritos, matriculados, "
            "matriculados_primer_curso",
            "2018–2024",
            "Excel (.xlsx)",
        ],
    )
    fill_table_row(
        table.rows[2],
        [
            "ICFES — Saber 3-5-9",
            "Scores de pruebas estandarizadas",
            "2018–2024",
            "CSV (.csv)",
        ],
    )
    fill_table_row(
        table.rows[3],
        [
            "PND (Plan Nacional de Desarrollo)",
            "Seguimiento de indicadores SINERGIA",
            "2018–2024",
            "CSV (.csv)",
        ],
    )

    doc.add_paragraph()

    doc.add_paragraph(
        "Los datos SNIES representan 7 categorías × 7 años = 49 archivos Excel. "
        "El pipeline descarga únicamente archivos con tamaño inferior a 15 MB para "
        "garantizar la viabilidad computacional. El indicador SINERGIA ID 91 (estudiantes "
        "nuevos en educación superior pública) es la variable de resultado principal "
        "alineada con la política evaluada."
    )
    doc.add_page_break()


def build_etl_pipeline(doc: Document) -> None:
    h = doc.add_heading("3. Pipeline ETL (Hito 2)", level=1)
    set_heading_color(h, COLOR_TITULO)

    doc.add_paragraph(
        "El pipeline ETL transforma los datos brutos en un esquema estrella PostgreSQL "
        "listo para el análisis estadístico. Está implementado como un orquestador "
        "de 11 pasos secuenciales con reporte de estado y manejo de errores."
    )

    h2 = doc.add_heading("3.1 Pasos del pipeline", level=2)
    set_heading_color(h2, COLOR_SECCION)

    steps = [
        ("1", "Ingesta de datos", "etl/ingest.py",
         "Descarga archivos desde Google Drive con autenticación OAuth2, retry "
         "exponencial y control de hashes MD5 para evitar re-importaciones."),
        ("2", "Carga a PostgreSQL (schema raw)", "scripts/create_db.py",
         "Detecta la fila de encabezado en cada Excel, normaliza nombres de columnas "
         "y usa COPY FROM STDIN para inserción masiva. Registra cada archivo "
         "importado en una tabla de log."),
        ("3", "Transformación y limpieza", "etl/transform.py",
         "Renombra columnas (snake_case), elimina columnas sin datos y registros "
         "vacíos, y limpia texto via pg_normalize_text(). Operaciones server-side en SQL."),
        ("4", "Normalización de datos", "scripts/normalize_data.py",
         "Aplica pg_normalize_text() a todas las columnas TEXT: elimina tildes, "
         "convierte a minúsculas, suprime espacios múltiples y strings vacíos."),
        ("5", "Creación de índices (raw)", "scripts/create_indexes.py",
         "Crea índices en los campos clave del schema raw (codigo_ies, ano, semestre, "
         "sexo, área) para acelerar las operaciones de unificación y JOIN."),
        ("6", "Unificación por año", "scripts/unify_by_year.py",
         "Consolida tablas por categoría (ej. matriculados_2018 ... matriculados_2024) "
         "en tablas unificadas en el schema unified con deduplicación."),
        ("7", "Creación de dimensiones", "scripts/create_dimensions.py",
         "Crea 7 tablas de dimensiones: dim_institucion, dim_geografia, dim_programa, "
         "dim_tiempo, dim_sexo, dim_nivel_formacion_docente, dim_dedicacion_docente."),
        ("8", "Creación de tablas de hechos", "scripts/create_facts.py",
         "Crea 3 fact tables con claves foráneas: fact_estudiantes (inscritos, admitidos, "
         "matriculados, primer_curso, graduados), fact_docentes, fact_administrativos."),
        ("9", "Verificación de calidad", "etl/quality.py",
         "Checks SQL via información_schema: tablas no vacías, duplicados, umbral de "
         "nulos (>50%), columnas mínimas. Genera reporte JSON."),
        ("10", "Diccionarios de datos", "etl/dictionary.py",
         "Perfila cada columna (count, nulos, min, max, mean, valores únicos) y genera "
         "diccionarios en JSON y Markdown."),
        ("11", "Export + subida a Google Drive", "etl/upload.py",
         "Exporta los 3 schemas (raw, unified, facts) con pg_dump y sube los archivos "
         ".sql a la carpeta compartida de Google Drive."),
    ]

    table = doc.add_table(rows=len(steps) + 1, cols=4)
    table.style = "Table Grid"
    add_table_header(table, ["Paso", "Nombre", "Módulo", "Descripción"])

    for i, (num, name, module, desc) in enumerate(steps):
        row = table.rows[i + 1]
        row.cells[0].text = num
        row.cells[1].text = name
        row.cells[2].text = module
        row.cells[3].text = desc

    doc.add_paragraph()

    h2 = doc.add_heading("3.2 Arquitectura del esquema estrella", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "El modelo de datos final (schema facts) implementa un esquema estrella con "
        "7 dimensiones y 3 tablas de hechos:"
    )

    h3 = doc.add_heading("Dimensiones", level=3)
    dims = [
        ("dim_institucion", "Instituciones de educación superior",
         "codigo_ies, nombre_ies, sector_ies (Oficial/Privada), caracter_ies, ..."),
        ("dim_geografia", "Ubicación geográfica",
         "codigo_departamento, nombre_departamento, codigo_municipio, nombre_municipio"),
        ("dim_programa", "Programas académicos",
         "codigo_snies_programa, nombre_programa, nivel_academico, nivel_formacion, area_conocimiento, ..."),
        ("dim_tiempo", "Periodos temporales",
         "ano (2018–2024), semestre (1 o 2), ano_semestre"),
        ("dim_sexo", "Género del estudiante/docente",
         "id_sexo, sexo (Masculino, Femenino, No binario, ...)"),
        ("dim_nivel_formacion_docente", "Nivel de formación del docente",
         "id_nivel, nivel_formacion_docente (Doctorado, Maestria, ...)"),
        ("dim_dedicacion_docente", "Tipo de dedicación del docente",
         "id_tiempo_dedicacion, tiempo_dedicacion, id_tipo_contrato, tipo_contrato"),
    ]

    table2 = doc.add_table(rows=len(dims) + 1, cols=3)
    table2.style = "Table Grid"
    add_table_header(table2, ["Tabla", "Descripción", "Columnas principales"])
    for i, (t, d, c) in enumerate(dims):
        table2.rows[i + 1].cells[0].text = t
        table2.rows[i + 1].cells[1].text = d
        table2.rows[i + 1].cells[2].text = c

    doc.add_paragraph()

    h3 = doc.add_heading("Tablas de hechos", level=3)
    facts = [
        ("fact_estudiantes",
         "Inscritos, admitidos, matriculados, primer_curso y graduados",
         "tipo_evento, institucion_id, programa_id, geografia_ies_id, "
         "geografia_programa_id, sexo_id, tiempo_id, cantidad"),
        ("fact_docentes",
         "Docentes por institución, dedicación y nivel de formación",
         "institucion_id, geografia_ies_id, sexo_id, nivel_formacion_docente_id, "
         "dedicacion_docente_id, tiempo_id, cantidad_docentes"),
        ("fact_administrativos",
         "Personal administrativo por institución y periodo",
         "institucion_id, geografia_ies_id, tiempo_id, auxiliar, tecnico, "
         "profesional, directivo, total"),
    ]

    table3 = doc.add_table(rows=len(facts) + 1, cols=3)
    table3.style = "Table Grid"
    add_table_header(table3, ["Tabla de hechos", "Descripción", "Columnas principales"])
    for i, (t, d, c) in enumerate(facts):
        table3.rows[i + 1].cells[0].text = t
        table3.rows[i + 1].cells[1].text = d
        table3.rows[i + 1].cells[2].text = c

    doc.add_paragraph()

    h2 = doc.add_heading("3.3 Características técnicas del ETL", level=2)
    set_heading_color(h2, COLOR_SECCION)

    caracteristicas = [
        ("Idempotencia",
         "Uso de hashes MD5 para evitar re-importaciones. Las tablas usan "
         "CREATE TABLE IF NOT EXISTS + TRUNCATE en lugar de DROP TABLE."),
        ("Rendimiento",
         "Inserciones masivas con COPY FROM STDIN (create_db) y execute_values "
         "(dimensiones, hechos). Transformaciones server-side en PostgreSQL via SQL."),
        ("Calidad de datos",
         "Checks SQL sin cargar datos en memoria: COUNT, COUNT DISTINCT, "
         "information_schema. Umbral configurable de nulos (50%)."),
        ("Seguridad",
         "Conexiones con psycopg2.sql.SQL + Identifier para prevenir SQL injection. "
         "Contenedor Docker con usuario no-root."),
        ("Resiliencia",
         "Descargas con retry + exponential backoff, archivos temporales atómicos "
         "(.tmp), conexiones con TCP keepalive."),
        ("Orquestación",
         "Pipeline local (etl/pipeline.py) y DAG de Airflow con 4 TaskGroups: "
         "ingestion, staging, star_schema, delivery."),
    ]

    table4 = doc.add_table(rows=len(caracteristicas) + 1, cols=2)
    table4.style = "Table Grid"
    add_table_header(table4, ["Característica", "Descripción"])
    for i, (c, d) in enumerate(caracteristicas):
        table4.rows[i + 1].cells[0].text = c
        table4.rows[i + 1].cells[1].text = d

    doc.add_page_break()


def build_metodologia(doc: Document) -> None:
    h = doc.add_heading("4. Metodología de Evaluación Causal (Hito 3)", level=1)
    set_heading_color(h, COLOR_TITULO)

    # ITS
    h2 = doc.add_heading("4.1 Series de Tiempo Interrumpidas (ITS)", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "El análisis ITS estima el efecto del cambio de gobierno sobre la matrícula "
        "en IES Oficiales usando el modelo de regresión segmentada:"
    )

    p = doc.add_paragraph()
    run = p.add_run(
        "Y_t = α₀ + α₁·t + α₂·D_t + α₃·(t − T₀)·D_t + ε_t"
    )
    run.font.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph(
        "donde:"
    )

    params = [
        ("t", "Índice de tiempo secuencial (t=1 corresponde a 2018-S1)"),
        ("D_t = 1 si t ≥ T₀", "Indicador post-intervención (T₀ = 2022-S2)"),
        ("α₀", "Nivel inicial (intercepto)"),
        ("α₁", "Tendencia pre-intervención por semestre"),
        ("α₂", "Cambio inmediato de nivel en el punto de quiebre"),
        ("α₃", "Cambio en la pendiente (tendencia) post-intervención"),
    ]

    table = doc.add_table(rows=len(params) + 1, cols=2)
    table.style = "Table Grid"
    add_table_header(table, ["Parámetro", "Interpretación"])
    for i, (p_, d) in enumerate(params):
        table.rows[i + 1].cells[0].text = p_
        table.rows[i + 1].cells[1].text = d

    doc.add_paragraph()
    doc.add_paragraph(
        "Los errores se estiman con corrección HAC (Newey-West, maxlags=2) para "
        "controlar la autocorrelación inherente a las series temporales. El punto de "
        "intervención T₀ = 2022-S2 corresponde al inicio efectivo del gobierno Petro "
        "en agosto de 2022. Se utilizan los semestres 2018-S1 a 2024-S2 (14 puntos "
        "temporales)."
    )

    h3 = doc.add_heading("Supuestos y pruebas de robustez (ITS)", level=3)
    supuestos = [
        ("Tendencia pre-2022 habría continuado sin cambio de gobierno",
         "Plausible pero no verificable directamente",
         "Prueba placebo: desplazar T₀ a 2021-S2 y 2020-S2"),
        ("No hay shocks simultáneos en T₀",
         "Parcialmente violado: efectos rezagados de pandemia aún presentes en 2022",
         "Incluir variable de contexto (tasa de desempleo DANE)"),
        ("Errores ε_t no autocorrelacionados",
         "Posible autocorrelación en series semestrales",
         "Errores Newey-West con lag=2; prueba Durbin-Watson"),
        ("Series comparables año a año",
         "Documentado en Hito 2 (cambios metodológicos SNIES)",
         "Sensibilidad: excluir años de transición metodológica"),
    ]

    table2 = doc.add_table(rows=len(supuestos) + 1, cols=3)
    table2.style = "Table Grid"
    add_table_header(table2, ["Supuesto", "Evaluación", "Prueba de robustez"])
    for i, (s, e, p_) in enumerate(supuestos):
        table2.rows[i + 1].cells[0].text = s
        table2.rows[i + 1].cells[1].text = e
        table2.rows[i + 1].cells[2].text = p_

    doc.add_paragraph()

    # DiD
    h2 = doc.add_heading("4.2 Diferencias en Diferencias (DiD)", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "El DiD compara la evolución de la matrícula entre IES Oficiales (grupo tratado) "
        "e IES Privadas (grupo control) antes y después de 2022-S2."
    )

    h3 = doc.add_heading("Modelo TWFE simplificado (datos agregados)", level=3)
    p = doc.add_paragraph()
    run = p.add_run(
        "Y_{st} = α + β₁·POST_t + β₂·OFICIAL_s + β₃·(POST_t × OFICIAL_s) + ε_{st}"
    )
    run.font.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph(
        "El estimador de interés es β₃: diferencia en el cambio de matrícula entre "
        "sector oficial y privado post-2022, usando errores robustos HC1."
    )

    h3 = doc.add_heading("Modelo panel TWFE (efectos fijos de IES)", level=3)
    p = doc.add_paragraph()
    run = p.add_run(
        "ln(Y_{it} + 1) = μ_i + γ_t + β·(POST_t × OFICIAL_i) + ε_{it}"
    )
    run.font.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph(
        "Donde μ_i son efectos fijos de IES y γ_t son efectos fijos de tiempo. "
        "El estimador β aproxima el cambio porcentual diferencial: "
        "(exp(β) − 1) × 100%. Se filtran IES con al menos 4 periodos de datos."
    )

    h3 = doc.add_heading("Supuesto de tendencias paralelas", level=3)
    doc.add_paragraph(
        "El supuesto central del DiD se verifica mediante:"
    )
    for txt in [
        "Gráfico de pre-tendencias: evolución de ln(matrícula) por sector, 2018–2022.",
        "Event study: regresión de la interacción por año pre-tratamiento; los coeficientes "
        "pre-2022 no deben diferir significativamente de cero.",
        "Prueba placebo sectorial: DiD entre subgrupos de privadas (élite vs. masivas); "
        "el estimador debe ser cercano a cero.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_paragraph()

    # Bootstrap
    h2 = doc.add_heading("4.3 Bootstrap e Intervalos de Confianza (Criterio F)", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "Para los estimadores ITS (α₂, α₃) y DiD (β₃) se calculan intervalos de "
        "confianza mediante bootstrap no paramétrico con B = 1,000 re-muestras:"
    )

    for txt in [
        "Método: bootstrap por bloques temporales (block bootstrap con bloque = 2 semestres) "
        "para preservar la autocorrelación de la serie.",
        "Se reporta el IC 95% por método de percentiles (2.5% y 97.5%).",
        "Seed fijo (42) para reproducibilidad.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    h3 = doc.add_heading("Análisis de sensibilidad", level=3)
    sensibilidades = [
        ("Punto de quiebre (T₀)",
         "2021-S2, 2022-S1, 2022-S2 (base), 2023-S1"),
        ("Muestra de IES en DiD",
         "Todas las IES / Solo universidades / Excluir IES < 500 estudiantes"),
        ("Forma funcional",
         "Niveles / Logaritmo / Primeras diferencias"),
        ("Grupo de control DiD",
         "Solo privadas / Solo privadas con mismo departamento"),
    ]

    table3 = doc.add_table(rows=len(sensibilidades) + 1, cols=2)
    table3.style = "Table Grid"
    add_table_header(table3, ["Dimensión de sensibilidad", "Variaciones analizadas"])
    for i, (d, v) in enumerate(sensibilidades):
        table3.rows[i + 1].cells[0].text = d
        table3.rows[i + 1].cells[1].text = v

    doc.add_paragraph()

    h2 = doc.add_heading("4.4 Fuentes de incertidumbre", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph("Se distinguen tres fuentes de incertidumbre cuantificadas:")

    fuentes = [
        ("Incertidumbre de medición",
         "Rezagos de reporte SNIES, revisiones retroactivas",
         "IC bootstrap sobre los datos observados"),
        ("Incertidumbre de modelo",
         "Especificación lineal vs. cuadrática, lag de efectos",
         "Comparar AIC/BIC entre especificaciones alternativas"),
        ("Incertidumbre de identificación",
         "Validez del supuesto de tendencias paralelas",
         "Gráfico y prueba formal de pre-tendencias"),
    ]

    table4 = doc.add_table(rows=len(fuentes) + 1, cols=3)
    table4.style = "Table Grid"
    add_table_header(table4, ["Fuente", "Descripción", "Cuantificación"])
    for i, (f, d, c) in enumerate(fuentes):
        table4.rows[i + 1].cells[0].text = f
        table4.rows[i + 1].cells[1].text = d
        table4.rows[i + 1].cells[2].text = c

    doc.add_page_break()


def build_implementacion(doc: Document) -> None:
    h = doc.add_heading("5. Implementación Técnica", level=1)
    set_heading_color(h, COLOR_TITULO)

    doc.add_paragraph(
        "Todos los módulos están implementados en Python 3.12+ usando las siguientes "
        "bibliotecas principales: pandas, numpy, statsmodels, scipy, plotly, psycopg2, "
        "SQLAlchemy, streamlit."
    )

    h2 = doc.add_heading("5.1 Módulos de análisis (analysis/)", level=2)
    set_heading_color(h2, COLOR_SECCION)

    modulos = [
        ("analysis/queries.py",
         "Consultas SQL al star schema PostgreSQL. Retorna DataFrames pandas listos "
         "para análisis. Funciones: get_matricula_por_sector(), get_panel_ies(), "
         "get_embudo_estudiantil(), get_matricula_por_departamento(), "
         "get_docentes_por_sector()."),
        ("analysis/tendencias.py",
         "Análisis descriptivo: variación % anual e interperiodo. Genera gráficos "
         "de series temporales (Plotly) y resumen pre/post 2022 por sector."),
        ("analysis/its.py",
         "Modelo ITS con OLS + HAC (Newey-West). Calcula coeficientes α₀-α₃, "
         "prueba de Chow, contrafactual proyectado, placebos temporales, y genera "
         "gráfico observado vs. contrafactual."),
        ("analysis/did.py",
         "DiD agregado (sector × semestre) y panel TWFE (efectos fijos de IES y tiempo). "
         "Event study con prueba formal de pre-tendencias. Genera gráficos DiD y "
         "event study."),
        ("analysis/bootstrap.py",
         "Block bootstrap (bloque=2 semestres, N=1,000) para IC 95% de α₂, α₃ "
         "(ITS) y β₃ (DiD). Análisis de escenarios (base / optimista / adverso). "
         "Histogramas de distribución bootstrap."),
        ("analysis/runner.py",
         "Orquestador principal: ejecuta en orden tendencias → ITS → DiD → bootstrap → "
         "resumen ejecutivo JSON. Guarda todos los resultados en data/results/. "
         "Uso: python analysis/runner.py"),
    ]

    for nombre, desc in modulos:
        p = doc.add_paragraph()
        run_name = p.add_run(nombre + ": ")
        run_name.font.bold = True
        run_name.font.color.rgb = COLOR_SECCION
        p.add_run(desc)

    h2 = doc.add_heading("5.2 Dashboard Streamlit (dashboard/app.py)", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "Dashboard interactivo accesible en http://localhost:8501 con 5 pestañas:"
    )

    tabs = [
        ("Tendencias", "Serie temporal de matrícula por sector con línea de intervención "
         "(2022-S2), variación % anual y métricas pre/post."),
        ("ITS", "Gráfico observado vs. contrafactual, tabla de coeficientes (α₀-α₃) "
         "con IC 95%, prueba de Chow, placebos temporales."),
        ("DiD", "Estimador β₃ con IC 95%, tabla 2×2 (medias por sector y periodo), "
         "gráfico DiD, resultados panel TWFE, event study de pre-tendencias."),
        ("Bootstrap", "IC 95% bootstrap para α₂ y β₃, gráfico de escenarios "
         "(observado / base / optimista / adverso)."),
        ("Resumen ejecutivo", "Hallazgos principales cargados desde JSON, interpretación "
         "técnicamente neutral con tabla de tipos de afirmación."),
    ]

    table = doc.add_table(rows=len(tabs) + 1, cols=2)
    table.style = "Table Grid"
    add_table_header(table, ["Pestaña", "Contenido"])
    for i, (t, d) in enumerate(tabs):
        table.rows[i + 1].cells[0].text = t
        table.rows[i + 1].cells[1].text = d

    doc.add_paragraph()
    doc.add_paragraph(
        "El dashboard incluye un botón 'Re-ejecutar análisis' que invoca "
        "analysis/runner.py directamente desde la interfaz, actualizando todos los "
        "resultados sin necesidad de terminal."
    )

    h2 = doc.add_heading("5.3 Notebook Jupyter (notebooks/hito3_analisis.ipynb)", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "El notebook documenta el análisis paso a paso con celdas ejecutables:"
    )

    celdas = [
        "Carga de datos desde PostgreSQL via SQLAlchemy",
        "Análisis de tendencias y embudo estudiantil (inscripción → admisión → matrícula)",
        "ITS: coeficientes, bondad de ajuste, prueba de Chow, placebos",
        "DiD: estimador agregado, panel TWFE, event study",
        "Bootstrap: distribución, IC 95%, análisis de escenarios",
        "Análisis de sensibilidad al punto de quiebre y muestra de IES",
        "Resumen de hallazgos con interpretación técnicamente neutral",
    ]

    for c in celdas:
        doc.add_paragraph(c, style="List Number")

    doc.add_page_break()


def build_resultados(doc: Document) -> None:
    h = doc.add_heading("6. Resultados del Análisis", level=1)
    set_heading_color(h, COLOR_TITULO)

    doc.add_paragraph(
        "Los resultados del análisis se generan ejecutando el orquestador principal:"
    )

    p = doc.add_paragraph()
    run = p.add_run("    uv run python analysis/runner.py")
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    p.paragraph_format.left_indent = Inches(0.4)

    doc.add_paragraph(
        "Los resultados se guardan en la carpeta data/results/ con la siguiente estructura:"
    )

    resultados = [
        ("tendencias_*.csv", "Series temporales por sector y tipo de evento"),
        ("resumen_pre_post_*.csv", "Medias y cambio % pre/post 2022 por sector"),
        ("its_*.json", "Coeficientes ITS, bondad de ajuste, Chow, placebos"),
        ("its_datos_*.csv", "Serie temporal con contrafactual y efecto estimado"),
        ("did_agregado_*.json", "Estimador β₃, IC 95%, medias por sector"),
        ("did_panel_*.json", "Estimador panel TWFE, N IES, N obs"),
        ("event_study_*.csv", "Coeficientes y IC por semestre para pre-tendencias"),
        ("bootstrap_*.json", "IC 95% bootstrap para α₂ y β₃, escenarios"),
        ("resumen_ejecutivo_hito3.json", "Consolidado de hallazgos principales"),
        ("plots/*.html", "Gráficos interactivos Plotly (ITS, DiD, event study, bootstrap)"),
    ]

    table = doc.add_table(rows=len(resultados) + 1, cols=2)
    table.style = "Table Grid"
    add_table_header(table, ["Archivo", "Contenido"])
    for i, (f, d) in enumerate(resultados):
        table.rows[i + 1].cells[0].text = f
        table.rows[i + 1].cells[1].text = d

    doc.add_paragraph()

    h2 = doc.add_heading("6.1 Resultados ITS — Interpretación esperada", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "Para el sector Oficial (IES públicas), el modelo ITS estima:"
    )

    for txt in [
        "α₁ (tendencia pre-2022): crecimiento o decrecimiento semestral de la matrícula "
        "en el período 2018-S1 a 2022-S1, ajustando por tendencia post-pandemia.",
        "α₂ (cambio de nivel en 2022-S2): cambio inmediato en el número de estudiantes "
        "matriculados al inicio del gobierno Petro, independiente de la tendencia.",
        "α₃ (cambio de tendencia post-2022): si la pendiente de crecimiento cambió "
        "tras la implementación de políticas de gratuidad.",
        "El contrafactual proyecta qué habría ocurrido si la tendencia pre-2022 "
        "hubiera continuado, permitiendo estimar el efecto neto por semestre.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    h2 = doc.add_heading("6.2 Resultados DiD — Interpretación esperada", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "El estimador DiD β₃ cuantifica si las IES Oficiales crecieron diferencialmente "
        "más (o menos) que las privadas tras 2022-S2:"
    )

    for txt in [
        "β₃ > 0: las IES Oficiales ganaron más matrícula que las privadas (consistente "
        "con el efecto esperado de la política de gratuidad).",
        "β₃ ≈ 0: no hay diferencial entre sectores (la política no habría modificado "
        "la composición del crecimiento).",
        "β₃ < 0: las IES privadas crecieron más que las oficiales (inconsistente con "
        "la hipótesis del efecto de gratuidad).",
        "La significatividad estadística (p < 0.05) determina si el diferencial supera "
        "el ruido muestral.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_paragraph(
        "NOTA IMPORTANTE: los estimadores son medidas de asociación, no de causalidad "
        "pura. Un estimador positivo y significativo es consistente con la hipótesis "
        "de la política, pero no descarta explicaciones alternativas (recuperación "
        "post-COVID, ciclo económico, cambios demográficos)."
    )

    doc.add_page_break()


def build_neutralidad(doc: Document) -> None:
    h = doc.add_heading("7. Declaración de Neutralidad Técnica (Criterio G)", level=1)
    set_heading_color(h, COLOR_TITULO)

    doc.add_paragraph(
        "Los resultados de este análisis se presentan bajo los siguientes principios "
        "de neutralidad técnica:"
    )

    principios = [
        ("Hechos medidos",
         "Estadísticas descriptivas y estimadores de modelos con sus incertidumbres. "
         "Ejemplo: 'La matrícula en IES Oficiales registró X% de cambio entre 2022-S2 y 2024-S2'"),
        ("Inferencias condicionadas",
         "Interpretaciones causales explícitamente condicionadas a supuestos declarados. "
         "Ejemplo: 'Bajo el supuesto de tendencias paralelas, el DiD sugiere un diferencial "
         "de Y estudiantes atribuible a la política'"),
        ("Límites explícitos",
         "No se extrapola más allá de lo que los datos y el diseño permiten afirmar. "
         "Ejemplo: 'No es posible descartar que la recuperación post-pandemia u otros "
         "factores expliquen parte del cambio'"),
    ]

    table = doc.add_table(rows=len(principios) + 1, cols=2)
    table.style = "Table Grid"
    add_table_header(table, ["Tipo de afirmación", "Descripción y ejemplo"])
    for i, (t, d) in enumerate(principios):
        table.rows[i + 1].cells[0].text = t
        table.rows[i + 1].cells[1].text = d

    doc.add_paragraph()

    h2 = doc.add_heading("7.1 Posibles lecturas alternativas", level=2)
    set_heading_color(h2, COLOR_SECCION)

    doc.add_paragraph(
        "Un crítico de la política podría argumentar que el crecimiento post-2022 en "
        "IES Oficiales es continuación de la tendencia de recuperación post-COVID, no "
        "de la política de gratuidad. Esta lectura es igualmente compatible con los "
        "datos si α₂ y α₃ no son estadísticamente significativos."
    )

    doc.add_paragraph(
        "Un defensor de la política podría señalar que el diferencial positivo "
        "Oficial − Privada (β₃ > 0) es consistente con una política focalizada en "
        "el sector público que no debería afectar al privado de la misma forma. "
        "Esta lectura es válida si el supuesto de tendencias paralelas se sostiene."
    )

    doc.add_paragraph(
        "La evidencia disponible (14 puntos temporales semestrales, sin variación "
        "aleatoria de tratamiento) no permite inclinar definitivamente la balanza "
        "entre estas lecturas. El análisis aporta una estimación rigurosa y "
        "técnicamente neutral de la asociación entre la política y la matrícula, "
        "sin afirmar causalidad."
    )

    doc.add_page_break()


def build_conclusiones(doc: Document) -> None:
    h = doc.add_heading("8. Conclusiones", level=1)
    set_heading_color(h, COLOR_TITULO)

    conclusiones = [
        "Pipeline ETL completo: Se construyó un sistema de extracción, transformación "
        "y carga que procesa 49+ archivos SNIES (2018–2024), los consolida en un "
        "esquema estrella PostgreSQL con 7 dimensiones y 3 tablas de hechos, y exporta "
        "automáticamente los datos a Google Drive.",
        "Metodología causal aprobada: Se implementaron dos métodos complementarios "
        "(ITS y DiD) con corrección de errores HAC, pruebas de validación (Chow, "
        "pre-tendencias, placebos) y análisis de sensibilidad a supuestos clave.",
        "Incertidumbre cuantificada: Los intervalos de confianza bootstrap (N=1,000, "
        "block bootstrap) permiten reportar la incertidumbre de los estimadores sin "
        "asumir distribución paramétrica, distinguiendo tres fuentes: medición, "
        "modelo e identificación.",
        "Dashboard funcional: La aplicación Streamlit integra todos los resultados "
        "en una interfaz interactiva con 5 pestañas, incluyendo re-ejecución del "
        "análisis en tiempo real.",
        "Reproducibilidad: El notebook Jupyter y el orquestador runner.py permiten "
        "reproducir todo el análisis desde cero con un único comando, partiendo de "
        "la base de datos PostgreSQL.",
        "Neutralidad técnica: Los resultados se presentan con declaraciones "
        "explícitas sobre el tipo de inferencia (hecho medido, inferencia condicionada, "
        "límite del análisis), sin tomar partido por ninguna posición política.",
    ]

    for i, c in enumerate(conclusiones, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.font.bold = True
        p.add_run(c)

    doc.add_page_break()


def build_bibliografia(doc: Document) -> None:
    h = doc.add_heading("9. Bibliografía", level=1)
    set_heading_color(h, COLOR_TITULO)

    refs = [
        "Bernal, R. & Peña, X. (2011). Guía práctica para la evaluación de impacto. "
        "Cap. 4: Diferencias en diferencias. CEDE - Universidad de los Andes.",
        "Cameron, A. & Trivedi, P. (2005). Microeconometrics: Methods and Applications. "
        "Cambridge University Press. Cap. 11: Bootstrap.",
        "DNP (2022). Plan Nacional de Desarrollo 2022–2026: Colombia Potencia Mundial "
        "de la Vida. Departamento Nacional de Planeación, Bogotá.",
        "MEN / SNIES (2018–2024). Bases de datos de matrícula, admisiones, inscripciones, "
        "graduados y docentes en educación superior. Ministerio de Educación Nacional.",
        "Newey, W. & West, K. (1987). A simple, positive semi-definite, heteroskedasticity "
        "and autocorrelation consistent covariance matrix. Econometrica, 55(3), 703–708.",
        "Shadish, W., Cook, T. & Campbell, D. (2002). Experimental and Quasi-Experimental "
        "Designs for Generalized Causal Inference. Cap. 7: Interrupted Time Series. "
        "Houghton Mifflin.",
        "Wagner, A., Soumerai, S., Zhang, F. & Ross-Degnan, D. (2002). Segmented "
        "regression analysis of interrupted time series studies in medication use "
        "research. Journal of Clinical Pharmacy and Therapeutics, 27, 299–309.",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


def build_anexos(doc: Document) -> None:
    h = doc.add_heading("10. Anexos Técnicos", level=1)
    set_heading_color(h, COLOR_TITULO)

    h2 = doc.add_heading("10.1 Estructura del repositorio", level=2)
    set_heading_color(h2, COLOR_SECCION)

    estructura = [
        "config/         Constantes, rutas, credenciales, fuentes de datos",
        "etl/            Pipeline ETL: ingest, transform, quality, dictionary, upload, pipeline",
        "analysis/       Módulos de análisis causal: queries, tendencias, its, did, bootstrap, runner",
        "scripts/        Scripts de base de datos: create_db, create_schemas, normalize,",
        "                create_indexes, unify_by_year, create_dimensions, create_facts,",
        "                generate_report (este script)",
        "dashboard/      Aplicación Streamlit (app.py)",
        "notebooks/      Jupyter notebook del análisis (hito3_analisis.ipynb)",
        "utils/          Utilidades: db, google_auth, logger, text, schema_helpers",
        "airflow/        DAG de Airflow para orquestación en producción",
        "data/           Datos generados (raw/, processed/, results/, exports/)",
        "docs/           Documentación: metodologia_hito3.md, star_schema_design.md",
    ]

    for e in estructura:
        p = doc.add_paragraph()
        run = p.add_run(e)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    h2 = doc.add_heading("10.2 Dependencias principales", level=2)
    set_heading_color(h2, COLOR_SECCION)

    deps = [
        ("pandas >= 3.0", "Manipulación de datos tabulares"),
        ("numpy >= 2.2", "Cálculos numéricos y bootstrap"),
        ("statsmodels >= 0.14", "OLS con corrección HAC, regresión segmentada"),
        ("scipy >= 1.15", "Tests estadísticos (Chow, pre-tendencias)"),
        ("plotly >= 6.5", "Gráficos interactivos"),
        ("psycopg2-binary >= 2.9", "Conexión a PostgreSQL"),
        ("sqlalchemy >= 2.0", "ORM y pool de conexiones"),
        ("streamlit >= 1.19", "Dashboard interactivo"),
        ("python-docx >= 1.1", "Generación del informe Word (este script)"),
        ("google-api-python-client >= 2.19", "Integración con Google Drive"),
        ("openpyxl >= 3.1", "Lectura de archivos Excel (.xlsx)"),
        ("jupyter >= 1.1", "Notebook de análisis"),
    ]

    table = doc.add_table(rows=len(deps) + 1, cols=2)
    table.style = "Table Grid"
    add_table_header(table, ["Paquete", "Propósito"])
    for i, (p_, d) in enumerate(deps):
        table.rows[i + 1].cells[0].text = p_
        table.rows[i + 1].cells[1].text = d

    doc.add_paragraph()

    h2 = doc.add_heading("10.3 Comandos de ejecución", level=2)
    set_heading_color(h2, COLOR_SECCION)

    comandos = [
        ("Instalar dependencias", "uv sync"),
        ("Autenticación Google", "uv run python reauth.py"),
        ("Pipeline ETL completo", "uv run python -m etl.pipeline"),
        ("Solo análisis (sin ingesta)", "uv run python -m etl.pipeline --skip-ingest"),
        ("Análisis Hito 3", "uv run python analysis/runner.py"),
        ("Dashboard", "streamlit run dashboard/app.py"),
        ("Notebook Jupyter", "uv run jupyter notebook notebooks/hito3_analisis.ipynb"),
        ("Generar este informe", "uv run python scripts/generate_report.py"),
        ("Docker (pipeline)", "docker compose run pipeline python -m etl.pipeline"),
        ("Docker (dashboard)", "docker compose up dashboard"),
    ]

    table2 = doc.add_table(rows=len(comandos) + 1, cols=2)
    table2.style = "Table Grid"
    add_table_header(table2, ["Acción", "Comando"])
    for i, (a, c) in enumerate(comandos):
        table2.rows[i + 1].cells[0].text = a
        table2.rows[i + 1].cells[1].text = c


# ──────────────────────────────────────────────────────────────
# Función principal
# ──────────────────────────────────────────────────────────────

def generate_report() -> Path:
    print("Generando informe técnico Hito 3...")

    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Estilo de párrafo normal
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    style.paragraph_format.space_after = Pt(6)

    # Construir secciones
    build_portada(doc)
    build_resumen_ejecutivo(doc)
    build_introduccion(doc)
    build_fuentes_datos(doc)
    build_etl_pipeline(doc)
    build_metodologia(doc)
    build_implementacion(doc)
    build_resultados(doc)
    build_neutralidad(doc)
    build_conclusiones(doc)
    build_bibliografia(doc)
    build_anexos(doc)

    doc.save(str(OUTPUT_PATH))
    print(f"Informe generado exitosamente: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    generate_report()
